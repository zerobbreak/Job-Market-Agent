import logging
import os
import re
import json
import time
import traceback
import pypdf  # type: ignore
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime
from config import Config
from agents import interview_prep_agent
from utils import AdvancedJobScraper, CVTailoringEngine
from services import job_store
from utils.pdf_generator import PDFGenerator
from utils.scraping import extract_skills_from_description
from utils.ai_retries import retry_ai_call
from services.matching_service import SemanticMatcher
from appwrite.services.tables_db import TablesDB  # type: ignore
from appwrite.services.databases import Databases  # type: ignore  # Still needed for schema operations
from appwrite.services.storage import Storage  # type: ignore
from appwrite.query import Query  # type: ignore
from appwrite.id import ID  # type: ignore
from appwrite.client import Client  # type: ignore
import threading
from dataclasses import asdict

logger = logging.getLogger(__name__)

# Global storage for pipelines and profiles (Cache-Aside pattern)
pipeline_store = {}
profile_store = {}
store_lock = threading.Lock()

# Initialize Matcher
matcher = SemanticMatcher()

def parse_profile(profile_output):
    """Clean and normalize profile data"""
    if isinstance(profile_output, dict):
        return {
            'name': profile_output.get('name', ''),
            'email': profile_output.get('email', ''),
            'phone': profile_output.get('phone', ''),
            'location': profile_output.get('location', ''),
            'skills': profile_output.get('skills', []),
            'experience_level': profile_output.get('experience_level', 'Entry Level'),
            'education': profile_output.get('education', []),
            'work_experience': profile_output.get('work_experience', []),
            'projects': profile_output.get('projects', []),
            'strengths': profile_output.get('strengths', []),
            'career_goals': profile_output.get('career_goals', ''),
            'links': profile_output.get('links', {'linkedin': '', 'github': '', 'portfolio': ''})
        }

    profile_data = {'skills': [], 'experience_level': '', 'education': '', 'strengths': [], 'career_goals': '', 'name': '', 'email': '', 'phone': '', 'location': '', 'links': {'linkedin': '', 'github': '', 'portfolio': ''}}
    try:
        profile_text = str(profile_output)
        if hasattr(profile_output, 'decode'):
            profile_text = profile_output.decode('utf-8', errors='ignore')

        json_str = profile_text
        if '```json' in profile_text:
            json_str = profile_text.split('```json')[1].split('```')[0]
        elif '```' in profile_text:
            json_str = profile_text.split('```')[1].split('```')[0]
        
        json_str = json_str.strip()
        try:
            data = json.loads(json_str)
            if isinstance(data, dict):
                return parse_profile(data)
        except Exception:
            pass

        # Regex fallback
        exp_match = re.search(r'(?:Experience Level|Experience|Years)[:\s]*([^\n]+)', profile_text, re.IGNORECASE)
        if exp_match: profile_data['experience_level'] = exp_match.group(1).strip()
        
        edu_match = re.search(r'(?:Education|Qualifications|Degrees)[:\s]*([\s\S]*?)(?=\n{2,}|\n[A-Z]|$)', profile_text, re.IGNORECASE)
        if edu_match: profile_data['education'] = edu_match.group(1).strip()
        
    except Exception as e:
        logger.error(f"Error parsing profile: {e}")
        
    return profile_data


def ensure_database_schema():
    """Ensure Appwrite schema exists."""
    try:
        api_key = Config.APPWRITE_API_KEY
        if not api_key:
            logger.warning("APPWRITE_API_KEY not found. Schema checks skipped.")
            return

        admin_client = Client()
        admin_client.set_endpoint(Config.APPWRITE_ENDPOINT)
        admin_client.set_project(Config.APPWRITE_PROJECT_ID)
        admin_client.set_key(api_key)
        # Schema operations still use the old Databases API
        # TablesDB is for data operations (list_rows, create_row, etc.)
        admin_db = Databases(admin_client)

        def _create_attr(db, db_id, coll_id, attr_id, size, required=False):
            try:
                # Schema operations use the old Databases API
                # Use keyword arguments for optional parameters to ensure correct method signature
                db.create_string_attribute(
                    db_id, 
                    coll_id, 
                    attr_id, 
                    size=size, 
                    required=required
                )
            except Exception as e:
                if '409' not in str(e) and 'already exists' not in str(e).lower():
                    logger.error(f"Failed to create attribute '{attr_id}': {e}")

        try:
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'education', 2000)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'experience_level', 255)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'career_goals', 2000)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'strengths', 2500)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'cv_hash', 64)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'name', 255)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'email', 255)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'phone', 50)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'location', 255)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'work_experience', 10000, False)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_PROFILES, 'projects', 10000, False)

            # Jobs Collection Schema
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_JOBS, 'status', 50, False)
            # Integer attribute
            try:
                admin_db.create_integer_attribute(
                    Config.DATABASE_ID, 
                    Config.COLLECTION_ID_JOBS, 
                    'progress', 
                    required=False, 
                    min=0, 
                    max=100, 
                    default=0
                )
            except Exception: pass

            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_JOBS, 'phase', 255, False)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_JOBS, 'job_data', 10000, False) # JSON string
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_JOBS, 'result', 10000, False) # JSON string
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_JOBS, 'template_type', 50, False)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_JOBS, 'user_id', 50, True)
            _create_attr(admin_db, Config.DATABASE_ID, Config.COLLECTION_ID_JOBS, 'error', 5000, False)
            
        except Exception as e:
            logger.error(f"Schema update error: {e}")
        
        logger.info("Database schema ensured (Basic Check)")
    except Exception as e:
        logger.error(f"ensure_database_schema error: {e}")

def _rehydrate_pipeline_from_profile(session_id: str, client) -> 'JobApplicationPipeline | None':
    try:
        logger.info(f"Attempting rehydration for user {session_id}")
        tablesDB = TablesDB(client)
        storage = Storage(client)
        
        existing_profiles = tablesDB.list_rows(
            Config.DATABASE_ID,
            Config.COLLECTION_ID_PROFILES,
            queries=[Query.equal('userId', session_id)]
        )

        if existing_profiles.get('total', 0) == 0:
            logger.info("No profile document found in DB")
            return None
        
        # TablesDB returns 'rows' instead of 'documents' in the new API
        rows = existing_profiles.get('rows', existing_profiles.get('documents', []))
        if not rows:
            logger.info("No profile rows found in DB")
            return None
            
        doc = rows[0]
        file_id = doc.get('cv_file_id')
        cv_text = doc.get('cv_text')
        
        logger.info(f"Rehydration: file_id={file_id}, cv_text length={len(cv_text) if cv_text else 0}")
        
        if cv_text:
            logger.info("Rehydrating from stored cv_text")
            pipeline = JobApplicationPipeline()
            pipeline.build_profile(cv_text)
            
            with store_lock:
                pipeline_store[session_id] = pipeline
                profile_store[session_id] = {
                    'profile_data': parse_profile(pipeline.profile),
                    'raw_profile': pipeline.profile,
                    'cv_filename': doc.get('cv_filename', 'CV'),
                    'cv_content': cv_text,
                    'file_id': file_id
                }
            return pipeline
        
        if file_id:
            logger.info(f"Downloading CV file {file_id} for rehydration")
            try:
                tmp_dir = Config.UPLOAD_FOLDER
                os.makedirs(tmp_dir, exist_ok=True)
                tmp_name = f"rehydrated_{session_id}_{file_id}.pdf"
                cv_path = os.path.join(tmp_dir, tmp_name)
                
                # Check if file exists and is valid size (>1KB)
                if not os.path.exists(cv_path) or os.path.getsize(cv_path) < 1024:
                    logger.info(f"Starting download of file {file_id} from bucket {Config.BUCKET_ID_CVS}...")
                    try:
                         data = storage.get_file_download(bucket_id=Config.BUCKET_ID_CVS, file_id=file_id)
                         with open(cv_path, 'wb') as f:
                             f.write(data)
                         logger.info(f"Download completed: {cv_path} ({os.path.getsize(cv_path)} bytes)")
                    except Exception as download_error:
                         logger.error(f"Download failed for file_id {file_id}: {download_error}")
                         import traceback
                         logger.error(traceback.format_exc())
                         return None
                else:
                     logger.info(f"Using cached file: {cv_path}")
                
                pipeline = JobApplicationPipeline(cv_path=cv_path)
                logger.info("Loading CV content...")
                cv_content = pipeline.load_cv()
                if not cv_content or len(cv_content) < 100:
                    logger.error(f"Loaded CV content is too short: {len(cv_content) if cv_content else 0} chars")
                    return None
                logger.info("Building profile from CV...")
                pipeline.build_profile(cv_content)
                logger.info("Profile built successfully.")
                
                with store_lock:
                    pipeline_store[session_id] = pipeline
                    profile_store[session_id] = {
                        'profile_data': parse_profile(pipeline.profile),
                        'raw_profile': pipeline.profile,
                        'cv_filename': doc.get('cv_filename', 'CV'),
                        'cv_content': cv_content,
                        'file_id': file_id
                    }
                logger.info(f"Rehydration successful: stored {len(cv_content)} chars in profile_store")
                return pipeline
            except Exception as e:
                logger.error(f"Download rehydration failed: {e}")
                import traceback
                logger.error(traceback.format_exc())
        else:
            logger.warning("No file_id found in profile document - cannot rehydrate from file")
                
        return None
    except Exception as e:
        logger.error(f"Rehydration overall failure: {e}")
        return None

class JobApplicationPipeline:
    def __init__(self, cv_path=None, output_dir='applications'):
        self.scraper = AdvancedJobScraper()
        self.cv_path = cv_path or Config.UPLOAD_FOLDER
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.cv_engine = None
        self.profile = None
        self.applications = []
        
    def _normalize_cv_text(self, text: str) -> str:
        """Enhanced text normalization for CV extraction - preserves structure"""
        try:
            if not text:
                return ""
            
            # Remove PDF ligatures and special character artifacts
            text = re.sub(r'\(cid:\d+\)', ' ', text)
            text = re.sub(r'\ufeff', '', text)  # Remove BOM
            
            # Fix hyphenated line breaks: devel-\nopment -> development
            # Handle both soft and hard hyphens
            text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)
            text = re.sub(r"(\w)\u00ad\s*\n\s*(\w)", r"\1\2", text)  # Soft hyphen
            text = re.sub(r"(\w)\u2011\s*\n\s*(\w)", r"\1\2", text)  # Non-breaking hyphen
            
            # Fix concatenated words: add space before capital letters after lowercase
            # e.g., "MongoDBCSS3React.js" -> "MongoDB CSS3 React.js"
            # BUT preserve common abbreviations and acronyms
            text = re.sub(r'([a-z])([A-Z][a-z])', r'\1 \2', text)
            
            # Fix concatenated words with dots: "Storage.MongoDB" -> "Storage, MongoDB"
            text = re.sub(r'\.([A-Z][a-z])', r', \1', text)
            
            # Fix concatenated words with closing paren: "Apps)Azure" -> "Apps, Azure"
            text = re.sub(r'\)([A-Z][a-z])', r'), \1', text)
            
            # PRESERVE SECTION HEADERS - Don't normalize common CV section headers
            # Common section headers (case-insensitive matching)
            section_headers = [
                r'PROFESSIONAL\s+PROFILE', r'PROFILE', r'SUMMARY', r'OBJECTIVE', 
                r'ABOUT\s+ME', r'CAREER\s+OBJECTIVE', r'PROFESSIONAL\s+SUMMARY',
                r'EXPERIENCE', r'WORK\s+EXPERIENCE', r'EMPLOYMENT', r'PROFESSIONAL\s+EXPERIENCE',
                r'EDUCATION', r'ACADEMIC', r'QUALIFICATIONS', r'STUDIES',
                r'SKILLS', r'TECHNICAL\s+SKILLS', r'CORE\s+COMPETENCIES', r'EXPERTISE',
                r'PROJECTS', r'KEY\s+PROJECTS', r'PORTFOLIO',
                r'CERTIFICATIONS', r'CERTIFICATES', r'AWARDS', r'ACHIEVEMENTS'
            ]
            
            # Mark section headers for preservation (we'll restore them later)
            section_markers = {}
            for i, pattern in enumerate(section_headers):
                matches = list(re.finditer(r'(?i)^\s*' + pattern + r'\s*:?\s*$', text, re.MULTILINE))
                for match in matches:
                    marker = f"__SECTION_MARKER_{i}_{match.start()}__"
                    section_markers[marker] = match.group(0).strip()
                    text = text[:match.start()] + marker + text[match.end():]
            
            # Fix spacing around punctuation (but preserve structure)
            text = re.sub(r'\s+([,.;:!?])', r'\1', text)
            text = re.sub(r'([,.;:!?])([A-Za-z])', r'\1 \2', text)
            
            # Remove excessive dots
            text = re.sub(r'\.{3,}', '...', text)
            
            # Fix multiple spaces and tabs (but preserve single line breaks for structure)
            text = re.sub(r'[ \t]+', ' ', text)
            
            # PRESERVE PARAGRAPH STRUCTURE - Keep single line breaks, limit only excessive ones
            # This helps preserve section boundaries and paragraph structure
            text = re.sub(r'\n{4,}', '\n\n\n', text)  # Max 3 line breaks
            # Don't collapse all line breaks - preserve paragraph structure
            
            # Remove common placeholder artifacts
            text = text.replace('\x0c', '\n')  # Form feed
            text = text.replace('\x00', '')    # Null bytes
            
            # Clean up bullet points (preserve them)
            text = re.sub(r'[•·▪▫]\s*', '• ', text)
            text = re.sub(r'(•\s*){2,}', '• ', text)
            
            # Restore section headers
            for marker, original_header in section_markers.items():
                text = text.replace(marker, original_header)
            
            # Clean up lines but preserve structure
            # Don't strip all lines completely - preserve empty lines that indicate structure
            lines = []
            for line in text.split('\n'):
                stripped = line.strip()
                if stripped:
                    lines.append(stripped)
                else:
                    # Preserve empty lines but limit consecutive ones
                    if not lines or lines[-1] != '':
                        lines.append('')
            
            text = '\n'.join(lines)
            
            # Final cleanup - but preserve meaningful structure
            text = text.strip()
            
        except Exception as e:
            logger.warning(f"Text normalization error: {e}")
            logger.debug(traceback.format_exc())
        
        return text

    def _extract_pdf_text(self, path: str) -> str:
        """Enhanced PDF text extraction with multiple fallback methods"""
        logger.info(f"Starting PDF text extraction from: {path}")
        extracted_texts = []
        
        # Method 1: pdfplumber (best for layout preservation)
        try:
            import pdfplumber  # type: ignore[import-untyped]
            text_parts = []
            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)
                logger.info(f"PDF has {page_count} page(s)")
                for i, page in enumerate(pdf.pages):
                    # Try multiple extraction strategies for better structure preservation
                    page_text = None
                    
                    # Strategy 1: Layout-based extraction (preserves structure)
                    try:
                        page_text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                    except:
                        pass
                    
                    # Strategy 2: Simple extraction if layout fails
                    if not page_text or len(page_text.strip()) < 50:
                        try:
                            page_text = page.extract_text()
                        except:
                            pass
                    
                    # Strategy 3: Extract by words and reconstruct (best for structure)
                    if not page_text or len(page_text.strip()) < 50:
                        try:
                            words = page.extract_words()
                            if words:
                                # Group words by vertical position to preserve paragraphs
                                lines = {}
                                for word in words:
                                    y = round(word['top'])
                                    if y not in lines:
                                        lines[y] = []
                                    lines[y].append((word['x0'], word['text']))
                                
                                # Sort by y position, then by x position within each line
                                sorted_lines = sorted(lines.items(), key=lambda x: -x[0])  # Top to bottom
                                page_lines = []
                                for y, words_in_line in sorted_lines:
                                    words_in_line.sort(key=lambda x: x[0])  # Left to right
                                    line_text = ' '.join([w[1] for w in words_in_line])
                                    if line_text.strip():
                                        page_lines.append(line_text)
                                
                                page_text = '\n'.join(page_lines)
                        except:
                            pass
                    
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Also try extracting tables if present
                    tables = page.extract_tables()
                    if tables:
                        logger.debug(f"Found {len(tables)} table(s) on page {i+1}")
                        for table in tables:
                            if table:
                                table_text = "\n".join([" | ".join([str(cell or "") for cell in row]) for row in table])
                                if table_text.strip():
                                    text_parts.append("\n[TABLE]\n" + table_text + "\n[/TABLE]\n")
            
            text = "\n".join(text_parts)
            
            # Detect and mark common CV sections for better parsing
            text = self._add_section_markers(text)
            
            if text and len(text.strip()) >= 200:
                preview = text[:500].replace('\n', ' ').strip()
                logger.info(f"Successfully extracted {len(text)} characters using pdfplumber (preview: {preview}...)")
                logger.debug(f"Full extracted text preview (first 1000 chars):\n{text[:1000]}")
                return text
            elif text:
                logger.warning(f"pdfplumber extracted only {len(text)} characters (below threshold)")
                extracted_texts.append(("pdfplumber", text))
        except ImportError:
            logger.warning("pdfplumber not available, trying other methods")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
            logger.debug(traceback.format_exc())
        
        # Method 2: pypdf (good for simple PDFs)
        try:
            reader = pypdf.PdfReader(path)
            text_parts = []
            for i, page in enumerate(reader.pages):
                # Try different extraction methods
                t = page.extract_text() or ""
                if not t or len(t.strip()) < 50:
                    # Try extracting with layout
                    t = page.extract_text(extraction_mode="layout") or ""
                text_parts.append(t)
            text = "\n".join(text_parts)
            # Detect and mark common CV sections
            text = self._add_section_markers(text)
            if text and len(text.strip()) >= 200:
                preview = text[:500].replace('\n', ' ').strip()
                logger.info(f"Successfully extracted {len(text)} characters using pypdf (preview: {preview}...)")
                logger.debug(f"Full extracted text preview (first 1000 chars):\n{text[:1000]}")
                return text
            elif text:
                logger.warning(f"pypdf extracted only {len(text)} characters (below threshold)")
                extracted_texts.append(("pypdf", text))
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")
            logger.debug(traceback.format_exc())
        
        # Method 3: pdfminer (fallback for complex PDFs)
        try:
            from pdfminer.high_level import extract_text as _pdfminer_extract_text  # type: ignore
            text = _pdfminer_extract_text(path) or ""
            # Detect and mark common CV sections
            text = self._add_section_markers(text)
            if text and len(text.strip()) >= 200:
                preview = text[:500].replace('\n', ' ').strip()
                logger.info(f"Successfully extracted {len(text)} characters using pdfminer (preview: {preview}...)")
                logger.debug(f"Full extracted text preview (first 1000 chars):\n{text[:1000]}")
                return text
            elif text:
                logger.warning(f"pdfminer extracted only {len(text)} characters (below threshold)")
                extracted_texts.append(("pdfminer", text))
        except ImportError:
            logger.warning("pdfminer not available")
        except Exception as e:
            logger.warning(f"pdfminer extraction unavailable/failed: {e}")
            logger.debug(traceback.format_exc())
        
        # Return the best result we have, even if short
        if extracted_texts:
            # Return the longest extracted text
            best_method, best_text = max(extracted_texts, key=lambda x: len(x[1]))
            # Detect and mark common CV sections
            best_text = self._add_section_markers(best_text)
            preview = best_text[:500].replace('\n', ' ').strip()
            logger.info(f"Using {best_method} extraction result ({len(best_text)} characters, preview: {preview}...)")
            logger.debug(f"Best extracted text preview (first 1000 chars):\n{best_text[:1000]}")
            return best_text
        
        logger.error(f"All PDF extraction methods failed for {path}")
        return ""
    
    def _add_section_markers(self, text: str) -> str:
        """Detect and mark common CV sections for better parsing"""
        if not text:
            return text
        
        # Common section header patterns (case-insensitive)
        section_patterns = [
            (r'(?i)^\s*(PROFESSIONAL\s+PROFILE|PROFILE|SUMMARY|OBJECTIVE|ABOUT\s+ME|CAREER\s+OBJECTIVE|PROFESSIONAL\s+SUMMARY)\s*:?\s*$', '[SECTION:PROFILE]'),
            (r'(?i)^\s*(EXPERIENCE|WORK\s+EXPERIENCE|EMPLOYMENT|PROFESSIONAL\s+EXPERIENCE|WORK\s+HISTORY)\s*:?\s*$', '[SECTION:EXPERIENCE]'),
            (r'(?i)^\s*(EDUCATION|ACADEMIC|QUALIFICATIONS|STUDIES|EDUCATIONAL\s+BACKGROUND)\s*:?\s*$', '[SECTION:EDUCATION]'),
            (r'(?i)^\s*(SKILLS|TECHNICAL\s+SKILLS|CORE\s+COMPETENCIES|EXPERTISE|TECHNICAL\s+EXPERTISE)\s*:?\s*$', '[SECTION:SKILLS]'),
            (r'(?i)^\s*(PROJECTS|KEY\s+PROJECTS|PORTFOLIO|PERSONAL\s+PROJECTS)\s*:?\s*$', '[SECTION:PROJECTS]'),
            (r'(?i)^\s*(CERTIFICATIONS|CERTIFICATES|AWARDS|ACHIEVEMENTS|HONORS)\s*:?\s*$', '[SECTION:CERTIFICATIONS]'),
        ]
        
        lines = text.split('\n')
        marked_lines = []
        
        for line in lines:
            marked = False
            for pattern, marker in section_patterns:
                if re.match(pattern, line):
                    marked_lines.append(marker)
                    marked_lines.append(line.strip())
                    marked = True
                    break
            if not marked:
                marked_lines.append(line)
        
        result = '\n'.join(marked_lines)
        
        # Log detected sections
        detected_sections = re.findall(r'\[SECTION:(\w+)\]', result)
        if detected_sections:
            logger.info(f"Detected CV sections: {', '.join(set(detected_sections))}")
        
        return result
    
    def _is_generic_phrase(self, text: str) -> bool:
        """Check if text contains only generic career phrases"""
        if not text or len(text.strip()) < 20:
            return False
        
        text_lower = text.lower().strip()
        generic_phrases = [
            "to leverage my skills",
            "seeking opportunities",
            "looking for",
            "to utilize my",
            "seeking a position",
            "looking to apply",
            "to contribute",
            "seeking employment",
            "to apply my",
            "seeking a role",
            "to leverage my skills in a challenging role",
            "seeking a challenging position",
            "looking for opportunities"
        ]
        
        # Check if text is primarily generic phrases
        generic_count = sum(1 for phrase in generic_phrases if phrase in text_lower)
        # If text is short and contains generic phrases, it's likely generic
        if len(text) < 50 and generic_count > 0:
            return True
        # If text is longer but mostly generic phrases
        if generic_count >= 2:
            return True
        # If text is very short and matches a generic phrase exactly
        if len(text) < 40:
            for phrase in generic_phrases:
                if phrase in text_lower and len(text_lower) - len(phrase) < 10:
                    return True
        
        return False
    
    def _get_career_goals(self, cv_data, cv_content: str) -> str:
        """Extract career goals with fallback logic"""
        # First, try to use AI-extracted professional_profile
        if cv_data and cv_data.professional_profile:
            profile = cv_data.professional_profile.strip()
            # Check if it's a generic fallback value using improved detection
            if not self._is_generic_phrase(profile):
                if len(profile) > 20:  # Meaningful length
                    logger.info(f"Using AI-extracted professional profile: {profile[:100]}...")
                    return profile
        
        # Fallback: Try rule-based extraction
        logger.warning("AI-extracted professional_profile is empty or generic, trying rule-based extraction")
        return self._get_career_goals_fallback(cv_content)
    
    def _get_career_goals_fallback(self, cv_content: str) -> str:
        """Rule-based extraction of career goals from CV content"""
        if not cv_content:
            logger.warning("No CV content available for fallback extraction")
            return ""
        
        try:
            from utils.cv_parser import CVParser
            parser = CVParser(raw_text=cv_content)
            fallback_profile = parser._extract_professional_profile_fallback(cv_content)
            
            if fallback_profile and len(fallback_profile.strip()) > 20:
                # Apply generic phrase filter to rule-based result
                if self._is_generic_phrase(fallback_profile):
                    logger.warning("Rule-based extraction found generic phrase, rejecting it")
                    return ""
                logger.info(f"Rule-based extraction found profile: {fallback_profile[:100]}...")
                return fallback_profile
            else:
                logger.warning("Rule-based extraction failed to find meaningful profile")
                return ""
        except Exception as e:
            logger.error(f"Error in rule-based profile extraction: {e}")
            logger.debug(traceback.format_exc())
            return ""

    def _extract_docx_text(self, path: str) -> str:
        """Enhanced DOCX text extraction with table and structure support"""
        try:
            import docx  # type: ignore
            doc = docx.Document(path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    text_parts.append(para_text)
            
            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join([cell for cell in row_cells if cell])
                    if row_text:
                        table_rows.append(row_text)
                if table_rows:
                    text_parts.append("\n".join(table_rows))
            
            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    header_text = "\n".join([para.text for para in section.header.paragraphs if para.text.strip()])
                    if header_text.strip():
                        text_parts.append(header_text)
                if section.footer:
                    footer_text = "\n".join([para.text for para in section.footer.paragraphs if para.text.strip()])
                    if footer_text.strip():
                        text_parts.append(footer_text)
            
            text = "\n".join(text_parts)
            if text and len(text.strip()) >= 50:
                logger.info(f"Successfully extracted {len(text)} characters from DOCX")
                return text
        except ImportError:
            logger.warning("python-docx not available, trying fallback")
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}")
        
        # Fallback: Extract from XML directly
        try:
            import zipfile
            with zipfile.ZipFile(path) as z:
                # Get main document
                xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
                
                # Better XML parsing - extract text nodes
                try:
                    root = ET.fromstring(xml)
                    # Define namespace
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    # Extract all text elements
                    text_elements = root.findall('.//w:t', ns)
                    text_parts = [elem.text for elem in text_elements if elem.text]
                    text = "\n".join([t.strip() for t in text_parts if t.strip()])
                    if text and len(text.strip()) >= 50:
                        logger.info(f"Successfully extracted {len(text)} characters from DOCX XML")
                        return text
                except ET.ParseError:
                    # Fallback to regex if XML parsing fails
                    pass
                
                # Regex fallback
                xml = re.sub(r"<[^>]+>", "\n", xml)
                xml = re.sub(r"&[a-z]+;", " ", xml)  # Replace entities
                lines = [l.strip() for l in xml.splitlines() if l.strip()]
                text = "\n".join(lines)
                if text and len(text.strip()) >= 50:
                    logger.info(f"Successfully extracted {len(text)} characters from DOCX (regex fallback)")
                    return text
        except Exception as e2:
            logger.warning(f"DOCX XML fallback failed: {e2}")
        
        logger.error(f"All DOCX extraction methods failed for {path}")
        return ""

    def _extract_doc_text(self, path: str) -> str:
        try:
            import textract  # type: ignore
            return textract.process(path).decode("utf-8", errors="ignore")
        except Exception as e:
            logger.warning(f"textract .doc extraction unavailable/failed: {e}")
        return ""

    def load_cv(self):
        """Load CV content from file with robust extraction"""
        logger.info(f"Loading CV from: {self.cv_path}")
        if not os.path.exists(self.cv_path):
            raise FileNotFoundError(f"CV not found at {self.cv_path}")

        ext = Path(self.cv_path).suffix.lower()
        logger.info(f"CV file extension: {ext}")
        content = ""
        if ext == ".pdf":
            content = self._extract_pdf_text(self.cv_path)
        elif ext == ".docx":
            content = self._extract_docx_text(self.cv_path)
        elif ext == ".doc":
            content = self._extract_doc_text(self.cv_path)
        else:
            try:
                with open(self.cv_path, "r", encoding="utf-8") as f:
                    content = f.read()
                logger.info(f"Read {len(content)} characters from text file")
            except Exception as e:
                logger.error(f"Unsupported CV format or read error: {e}")
                raise
        
        if not content or len(content.strip()) < 50:
            logger.warning(f"Extracted content is too short ({len(content)} chars), using placeholder")
            content = f"Extracted from {self.cv_path}: insufficient text for analysis."
        
        logger.info(f"Raw extracted content length: {len(content)} characters")
        content_before_norm = content[:500] if len(content) > 500 else content
        content = self._normalize_cv_text(content)
        logger.info(f"Normalized content length: {len(content)} characters")
        logger.debug(f"Content before normalization (first 500 chars): {content_before_norm}")
        logger.debug(f"Content after normalization (first 500 chars): {content[:500]}")
        return content
    
    @retry_ai_call
    def build_profile(self, cv_content):
        """Build student profile using Hybrid AI Parser"""
        try:
            from utils.cv_parser import CVParser
            
            # Always use raw_text to avoid PDF bytes issues
            parser = CVParser(raw_text=cv_content)
            if self.cv_path and self.cv_path.lower().endswith('.pdf'):
                # Set file_path for reference but use raw_text for parsing
                parser.file_path = self.cv_path
                
            cv_data = parser.parse()
            
            # Handle case where parsing returns None
            if not cv_data:
                logger.warning("CV parsing returned None, using fallback profile")
                raise ValueError("Failed to parse CV data")
            
            skills = []
            if cv_data.technical_skills:
                for cat, s_list in cv_data.technical_skills.items():
                    skills.extend(s_list)
            
            # Prepare structured data for frontend
            # Use model_dump() for Pydantic v2 or dict() for v1
            education_list = []
            experience_list = []
            projects_list = []
            
            try:
                # Try Pydantic v2 method first (model_dump)
                if hasattr(cv_data.education[0], 'model_dump') if cv_data.education else False:
                    education_list = [e.model_dump() for e in cv_data.education]
                    experience_list = [e.model_dump() for e in cv_data.work_experience]
                    projects_list = [p.model_dump() for p in cv_data.projects]
                # Try Pydantic v1 method (dict)
                elif hasattr(cv_data.education[0], 'dict') if cv_data.education else False:
                    education_list = [e.dict() for e in cv_data.education]
                    experience_list = [e.dict() for e in cv_data.work_experience]
                    projects_list = [p.dict() for p in cv_data.projects]
                else:
                    # Manual conversion fallback
                    for e in cv_data.education:
                        if hasattr(e, '__dict__'):
                            education_list.append(e.__dict__)
                        else:
                            education_list.append({k: getattr(e, k, None) for k in ['degree', 'institution', 'year', 'details']})
                    for e in cv_data.work_experience:
                        if hasattr(e, '__dict__'):
                            experience_list.append(e.__dict__)
                        else:
                            experience_list.append({k: getattr(e, k, None) for k in ['title', 'company', 'duration', 'responsibilities']})
                    for p in cv_data.projects:
                        if hasattr(p, '__dict__'):
                            projects_list.append(p.__dict__)
                        else:
                            projects_list.append({k: getattr(p, k, None) for k in ['name', 'description', 'technologies', 'details']})
            except Exception as e:
                # Final fallback: manually convert to dict
                logger.warning(f"Error converting Pydantic models to dict: {e}")
                education_list = []
                experience_list = []
                projects_list = []
                try:
                    for e in cv_data.education:
                        education_list.append({
                            'degree': getattr(e, 'degree', ''),
                            'institution': getattr(e, 'institution', ''),
                            'year': getattr(e, 'year', ''),
                            'details': getattr(e, 'details', [])
                        })
                    for e in cv_data.work_experience:
                        experience_list.append({
                            'title': getattr(e, 'title', ''),
                            'company': getattr(e, 'company', ''),
                            'duration': getattr(e, 'duration', ''),
                            'responsibilities': getattr(e, 'responsibilities', [])
                        })
                    for p in cv_data.projects:
                        projects_list.append({
                            'name': getattr(p, 'name', ''),
                            'description': getattr(p, 'description', ''),
                            'technologies': getattr(p, 'technologies', []),
                            'details': getattr(p, 'details', [])
                        })
                except Exception as e2:
                    logger.error(f"Complete failure to convert Pydantic models: {e2}")
                    # Return empty lists to prevent crash
                    education_list = []
                    experience_list = []
                    projects_list = []
            
            exp_level = "Entry Level"
            if cv_data.work_experience:
                years_of_exp = len(cv_data.work_experience) * 1.5
                if any('senior' in exp.title.lower() for exp in cv_data.work_experience) or years_of_exp > 5:
                    exp_level = "Senior"
                elif years_of_exp > 2:
                    exp_level = "Mid Level"
            
            self.profile = {
                "name": cv_data.contact_info.name or "Unknown",
                "email": cv_data.contact_info.email or "",
                "phone": cv_data.contact_info.phone or "",
                "location": cv_data.contact_info.address or "",
                "skills": list(set(skills)),
                "experience_level": exp_level,
                "education": education_list,
                "work_experience": experience_list,
                "projects": projects_list,
                "strengths": skills[:5],
                "career_goals": self._get_career_goals(cv_data, cv_content),
                "links": {
                    "linkedin": cv_data.contact_info.linkedin or "",
                    "github": cv_data.contact_info.github or "",
                    "portfolio": cv_data.contact_info.portfolio or ""
                }
            }
            
            self.cv_engine = CVTailoringEngine(cv_content, self.profile)
            return self.profile
            
        except Exception as e:
            logger.error(f"Error building profile: {e}")
            import traceback
            logger.error(traceback.format_exc())
            
            # Fallback: Extract basic info from text
            try:
                skills = extract_skills_from_description(cv_content)
                self.profile = {
                    "name": "Unknown",
                    "email": "",
                    "phone": "",
                    "location": "",
                    "skills": skills,
                    "experience_level": "Entry Level",
                    "education": [],
                    "work_experience": [],
                    "projects": [],
                    "strengths": skills[:5] if skills else [],
                    "career_goals": self._get_career_goals_fallback(cv_content),
                    "links": {"linkedin": "", "github": "", "portfolio": ""}
                }
                self.cv_engine = CVTailoringEngine(cv_content, self.profile)
                return self.profile
            except Exception as fallback_error:
                logger.error(f"Fallback profile creation also failed: {fallback_error}")
                # Return minimal profile to prevent complete failure
                self.profile = {
                    "name": "Unknown",
                    "email": "",
                    "phone": "",
                    "location": "",
                    "skills": [],
                    "experience_level": "Entry Level",
                    "education": [],
                    "work_experience": [],
                    "projects": [],
                    "strengths": [],
                    "career_goals": self._get_career_goals_fallback(cv_content),
                    "links": {"linkedin": "", "github": "", "portfolio": ""}
                }
                try:
                    self.cv_engine = CVTailoringEngine(cv_content, self.profile)
                except:
                    pass
                return self.profile

    def search_jobs(self, query, location, max_results, use_cache: bool = True):
        """Search for jobs using the scraper.

        Args:
            query: Search query string.
            location: Location string.
            max_results: Maximum number of jobs to return.
            use_cache: Whether to allow the scraper to use its on-disk cache.
                       For force-refresh flows (e.g. Live Feed), this should be False
                       so that we always hit the live job boards.
        """
        try:
            return self.scraper.scrape_jobs(
                site_name=['linkedin', 'indeed'],
                search_term=query,
                location=location,
                results_wanted=max_results,
                country_indeed=location,
                use_cache=use_cache,
            )
        except Exception as e:
            logger.error(f"Error searching jobs: {e}")
            return []

    def generate_application_package(self, job, template_type=None):
        """Generate optimized CV (PDF) and cover letter for a specific job"""
        job_title = job.get('title', 'Unknown Position')
        company = job.get('company', 'Unknown Company')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        app_dir_name = f"{company.replace(' ', '_').replace('/', '_')}_{job_title.replace(' ', '_').replace('/', '_')}_{timestamp}"
        app_dir = self.output_dir / app_dir_name
        app_dir.mkdir(parents=True, exist_ok=True)

        try:
            cv_content, ats_analysis = self.cv_engine.generate_tailored_cv(job, template_type)
            if not self.cv_engine.cv_versions:
                 logger.error("No CV versions generated")
                 return None
            version_id = list(self.cv_engine.cv_versions.keys())[-1]
            cv_data = self.cv_engine.get_cv_version(version_id) or {}

            pdf_path = self.cv_engine.export_cv(version_id, format='pdf', output_dir=str(app_dir))
            final_cv_path = app_dir / 'cv.pdf'
            try:
                os.replace(pdf_path, final_cv_path)
            except Exception:
                final_cv_path = Path(pdf_path)

            cover_letter_result = self.cv_engine.generate_cover_letter(job, tailored_cv=cv_content, output_dir=str(app_dir))
            if isinstance(cover_letter_result, str) and cover_letter_result.endswith('.pdf') and os.path.exists(cover_letter_result):
                cl_temp = Path(cover_letter_result)
                final_cl_path = app_dir / 'cover_letter.pdf'
                try:
                    os.replace(cl_temp, final_cl_path)
                except Exception:
                    final_cl_path = cl_temp
            else:
                final_cl_path = app_dir / 'cover_letter.txt'
                with open(final_cl_path, 'w', encoding='utf-8') as f:
                    f.write(cover_letter_result if isinstance(cover_letter_result, str) else str(cover_letter_result))

            metadata = {
                'job': {
                    'title': job_title,
                    'company': company,
                    'location': job.get('location', ''),
                    'url': job.get('url', '')
                },
                'cv': {
                    'version_id': version_id,
                    'ats_analysis': ats_analysis,
                    'ats_score': cv_data.get('ats_score'),
                    'job_keywords': cv_data.get('job_keywords')
                },
                'relevance_score': job.get('relevance_score'),
                'paths': {
                    'cv': str(final_cv_path),
                    'cover_letter': str(final_cl_path)
                }
            }
            metadata_path = app_dir / 'metadata.json'
            with open(metadata_path, 'w', encoding='utf-8') as mf:
                json.dump(metadata, mf, ensure_ascii=False, indent=2, default=str)

            app_id = job_store.save_application(
                job_data=job,
                cv_path=str(final_cv_path),
                cover_letter_path=str(final_cl_path),
                ats_score=cv_data.get('ats_score'),
                metadata_path=str(metadata_path),
                app_dir=str(app_dir)
            )

            return {
                'job_title': job_title,
                'company': company,
                'cv_path': str(final_cv_path),
                'cover_letter_path': str(final_cl_path),
                'app_dir': str(app_dir),
                'metadata_path': str(metadata_path),
                'ats_score': cv_data.get('ats_score'),
                'ats_analysis': ats_analysis,
                'app_id': app_id
            }

        except Exception as e:
            logger.error(f"Error generating application: {e}")
            return None

    @retry_ai_call
    def prepare_interview(self, job, output_dir=None):
        """Generate interview preparation materials for a job"""
        job_title = job.get('title', 'Unknown Position')
        company = job.get('company', 'Unknown Company')
        
        try:
            response = interview_prep_agent.run(f"""
            Prepare comprehensive interview materials for this position:
            Job: {job_title}
            Company: {company}
            Description: {job.get('description', 'Not provided')}
            """)
            
            out_dir = Path(output_dir) if output_dir else self.output_dir
            out_dir.mkdir(parents=True, exist_ok=True)
            
            prep_path = out_dir / 'interview_prep.pdf'
            try:
                generator = PDFGenerator()
                header = {'title': f"Interview Prep: {job_title}", 'company': company, 'date': datetime.now().strftime('%Y-%m-%d')}
                success = generator.generate_pdf(markdown_content=response.content, output_path=str(prep_path), template_name='modern', header=header)
                if success:
                    return str(prep_path)
            except Exception:
                pass
                
            txt_path = out_dir / 'interview_prep.txt'
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(response.content)
            return str(txt_path)
            
            return str(txt_path)
            
        except Exception as e:
            logger.error(f"Error preparing interview materials: {e}")
            return None

    def run(self, query, location, max_applications, template=None):
        """Run the full pipeline"""
        print("\n" + "=" * 80)
        print("🚀 JOB APPLICATION PIPELINE STARTED")
        print("=" * 80)
        
        cv_content = self.load_cv()
        if not self.build_profile(cv_content):
            return
            
        jobs = self.search_jobs(query, location, max_applications)
        if not jobs:
            return
            
        print(f"\n📝 Generating applications for top {len(jobs)} jobs...")
        
        for i, job in enumerate(jobs):
            print(f"\n--- Application {i+1}/{len(jobs)} ---")
            app_result = self.generate_application_package(job, template_type=template)
            
            if app_result:
                self.applications.append(app_result)
                self.prepare_interview(job, output_dir=app_result['app_dir'])
                
            if i < len(jobs) - 1:
                time.sleep(2)

    def print_summary(self):
        """Print pipeline summary"""
        print("\n" + "=" * 80)
        print("📊 PIPELINE SUMMARY")
        print("=" * 80)
        print(f"✓ Applications generated: {len(self.applications)}")
        print(f"✓ Output directory: {self.output_dir.absolute()}")
        
        if self.applications:
            print("\n📋 Generated Applications:")
            for i, app in enumerate(self.applications, 1):
                print(f"   {i}. {app['job_title']} at {app['company']}")
                print(f"      CV: {os.path.basename(app['cv_path'])}")
        
        print("\n" + "=" * 80)
        print("✅ PIPELINE COMPLETED")
        print("=" * 80)
